#!/usr/bin/env python3
"""Generate IEEE MA-RAG status report as DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUTPUT = Path(__file__).resolve().parent.parent / "docs" / "IEEE_STATUS_REPORT_JUNE2026.docx"
OUTPUT_ROOT = Path(__file__).resolve().parent.parent / "IEEE_STATUS_REPORT_JUNE2026.docx"


def set_default_font(document: Document, name: str = "Calibri", size: int = 11) -> None:
    style = document.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)


def add_title(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def add_subtitle(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_para(document: Document, text: str, bold: bool = False) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            table.rows[row_idx].cells[col_idx].text = value
    document.add_paragraph()


def build_report() -> Document:
    doc = Document()
    set_default_font(doc)

    add_title(doc, "IEEE Talent Meets AI — MA-RAG Project Status Report")
    add_subtitle(doc, "Technical Status Update · June 27, 2026")
    doc.add_paragraph()

    meta = [
        ("Report type", "Technical status update"),
        ("Project", "Multi-Agent Retrieval-Augmented Generation (MA-RAG)"),
        ("Task force", "IEEE Talent Meets AI"),
        ("Research basis", "MA-RAG: Multi-Agent RAG via Collaborative Chain-of-Thought Reasoning (arXiv:2505.20096)"),
        ("Prepared by", "Bhargav Boyapati, Volunteer Researcher"),
        ("Technical lead", "Chandra Shekar Konda, AI Technical Director"),
        ("Report date", "June 27, 2026"),
        ("Reporting period", "Project inception through June 27, 2026"),
        ("Repository", "MA-RAG — branch feature/track-b-agentic"),
        ("Project path", "/Users/bhargavboyapati/Projects/MA-RAG"),
    ]
    add_table(doc, ["Field", "Value"], meta)

    add_heading(doc, "1. Executive Summary", 1)
    add_para(
        doc,
        "This report documents the status of the MA-RAG research prototype through June 27, 2026. "
        "The project has progressed from a Phase 0 local RAG baseline to a demonstrable multi-agent "
        "system with eight specialized agents, an orchestrated workflow runtime, on-premises SLM "
        "inference, auditable evidence logging, API ingress with live agent tracing, corporate "
        "document ingestion, and passing automated validation.",
    )
    add_table(
        doc,
        ["Check", "Result (June 27, 2026)"],
        [
            ["Unit & integration tests (pytest)", "39 / 39 passed"],
            ["Knowledge-base regression (eval_kb.py)", "8 / 8 passed"],
            ["Active LLM", "Ollama llama3.2:3b (on-prem SLM)"],
            ["Local FAISS index", "Ready"],
            ["Web UI", "Wireframe only — build deferred"],
            ["Production cloud deployment", "Not started"],
        ],
    )
    add_para(
        doc,
        "Bottom line: The prototype is operationally complete for research demonstration on a local "
        "workstation. It aligns with the enterprise five-plane multi-agent RAG architecture in design. "
        "Production deployment (web UI, SSO, cloud services, encrypted confidential storage) remains a follow-on phase.",
        bold=True,
    )

    add_heading(doc, "1.1 Impact Summary — Used · Gained · Reduced (XYZ Format)", 2)
    add_para(
        doc,
        "The table below summarizes major technical choices (what we used), measurable outcomes "
        "(what we gained), and trade-offs or eliminations (what we reduced) through June 27, 2026.",
    )
    add_table(
        doc,
        ["What we USED (X)", "What we GAINED (Y)", "What we REDUCED (Z)"],
        [
            [
                "On-prem SLM — Ollama llama3.2:3b",
                "Local inference; no per-question API cost; data stays on machine during Q&A",
                "Dependence on cloud GPT; external data egress for LLM reasoning",
            ],
            [
                "Eight-agent MA-RAG workflow (Router → Critic)",
                "Explainable answers; auditable steps; aligns with IEEE / enterprise agentic RAG design",
                "Single black-box LLM prompt; unverifiable answer paths",
            ],
            [
                "FAISS + local docs/ folder (no SQL DB)",
                "Zero database ops; fast laptop prototype; simple rebuild via ingest.py",
                "Infrastructure cost and setup time for managed vector databases",
            ],
            [
                "Typed Pydantic agent contracts (messages.py)",
                "Stable agent I/O; same shapes for in-process and future remote A2A workers",
                "Fragile ad-hoc dict passing between workflow steps",
            ],
            [
                "JSONL evidence ledger + A2A journal per run",
                "Compliance-ready audit trail; Run ID links question → retrieval → agents → answer",
                "“Trust me” answers with no trace of how they were produced",
            ],
            [
                "SLM guardrails (slm_helpers.py, fast mode)",
                "Reliable demos on 3B model; acceptable latency in interactive chat",
                "Failed runs from weak JSON/planning on small model (without heuristics)",
            ],
            [
                "Corporate ingest (Word, Excel, PPT, OCR/password PDF)",
                "Real employer documents usable without manual conversion to .txt",
                "Friction of “convert everything to markdown first”",
            ],
            [
                "FastAPI + SSE (/ask/stream)",
                "API-ready ingress; live agent trace for UI and integrations",
                "Terminal-only visibility of agent pipeline",
            ],
            [
                "Consolidated docs/ + wiki excluded from index",
                "One folder for users; project KB not drowned by setup pages in search",
                "Confusion from multiple sample_doc folders and polluted retrieval",
            ],
            [
                "Automated tests (39) + KB eval (8/8)",
                "Regression safety after each change; objective proof for stakeholders",
                "Manual-only validation; silent regressions on demo questions",
            ],
            [
                "Interactive ask.py chat (default)",
                "Simple user UX: python ask.py → type question",
                "Long CLI flags (--agentic --stream) on every question",
            ],
        ],
    )
    add_heading(doc, "1.2 Overall XYZ Summary", 2)
    add_table(
        doc,
        ["Dimension", "Summary"],
        [
            [
                "USED",
                "MA-RAG multi-agent architecture + Ollama SLM + FAISS + FastAPI + JSONL audit + corporate document pipeline",
            ],
            [
                "GAINED",
                "End-to-end document Q&A on laptop; 8/8 KB eval; 39/39 tests; auditable agentic RAG aligned to enterprise five-plane design",
            ],
            [
                "REDUCED",
                "Cloud LLM cost (active path); ops complexity (no DB server); demo failures on SLM; user friction for ingest/ask; unexplainable answers",
            ],
        ],
    )

    add_heading(doc, "2. Project Background", 1)
    add_heading(doc, "2.1 Purpose", 2)
    add_para(
        doc,
        "The IEEE Talent Meets AI task force is evaluating how multi-agent RAG can move from academic "
        "research into practical, governable, enterprise-style systems. MA-RAG decomposes question "
        "answering into cooperating agent roles rather than a single LLM call.",
    )
    add_heading(doc, "2.2 Objectives", 2)
    add_table(
        doc,
        ["ID", "Objective", "Status"],
        [
            ["O1", "Local MA-RAG pipeline (ingest + ask)", "Complete"],
            ["O2", "Eight explicit agents with typed contracts", "Complete"],
            ["O3", "On-premises SLM (employer priority)", "Complete"],
            ["O4", "Per-run audit (evidence + A2A journal)", "Complete"],
            ["O5", "KB regression validation", "Complete (8/8)"],
            ["O6", "Enterprise architecture alignment", "Partial"],
            ["O7", "Web user interface", "Wireframe only — deferred"],
        ],
    )

    add_heading(doc, "3. Scope", 1)
    add_heading(doc, "3.1 In Scope (Delivered)", 2)
    add_bullets(
        doc,
        [
            "Local document ingestion and FAISS indexing",
            "Eight-agent workflow orchestration",
            "Ollama SLM integration with quality guardrails",
            "JSONL evidence ledger and A2A message journal",
            "FastAPI ingress and SSE live agent trace",
            "Corporate formats: Word, Excel, PowerPoint, OCR/password PDFs",
            "Interactive CLI chat (python ask.py)",
            "Automated test suite and KB eval harness",
        ],
    )
    add_heading(doc, "3.2 Out of Scope (Current Phase)", 2)
    add_bullets(
        doc,
        [
            "Production cloud deployment",
            "Web UI implementation",
            "SSO / IAM / role-based document access",
            "CLARIFY / ESCALATE workflow branches",
            "Enterprise connectors (SharePoint, Confluence)",
            "Encrypted confidential index at rest",
        ],
    )

    add_heading(doc, "4. Technical Stack", 1)
    add_table(
        doc,
        ["Layer", "Technology", "Notes"],
        [
            ["LLM (active)", "Ollama llama3.2:3b", "On-prem SLM; default for all LLM agents"],
            ["LLM (fallback)", "OpenAI gpt-4o-mini", "Configured but inactive unless switched"],
            ["Retrieval", "FAISS + chunks.jsonl", "CPU-friendly; hashing embeddings default"],
            ["Orchestration", "WorkflowEngine", "Eight-step happy path"],
            ["Contracts", "Pydantic (messages.py)", "Stable agent I/O shapes"],
            ["API", "FastAPI + Uvicorn", "/ask, /ask/stream, /health"],
            ["Audit", "JSONL per run", "Evidence ledger + A2A journal"],
            ["Language", "Python 3.9", "Virtual env .venv"],
        ],
    )

    add_heading(doc, "5. Architecture", 1)
    add_heading(doc, "5.1 Workflow Pipeline", 2)
    add_para(
        doc,
        "route → init_plan → retrieve → evidence_check → context_build → generate → finalize → verify",
    )
    add_heading(doc, "5.2 Eight Agents", 2)
    add_table(
        doc,
        ["Agent", "File", "Role", "Uses SLM?"],
        [
            ["Router", "agents/router_agent.py", "Simple vs multi-hop triage", "No"],
            ["Planner", "agents/planner_agent.py", "Decompose into plan steps", "Yes"],
            ["Retrieval", "agents/retrieval_agent.py", "FAISS top-k search", "No"],
            ["Evidence Curator", "agents/evidence_curator_agent.py", "Evidence sufficiency", "Yes"],
            ["Step Definer", "agents/step_definer_agent.py", "Per-step task definition", "Yes"],
            ["RAG Step", "agents/rag_step_agent.py", "Grounded answer generation", "Yes"],
            ["Summarizer", "agents/summarizer_agent.py", "Merge step answers", "Yes"],
            ["Critic", "agents/critic_agent.py", "Verify faithfulness", "Yes"],
        ],
    )
    add_para(doc, "Orchestrator: src/workflow/engine.py")
    add_para(doc, "Contracts: src/contracts/messages.py")

    add_heading(doc, "5.3 Enterprise Five-Plane Alignment", 2)
    add_table(
        doc,
        ["Plane", "Target", "Prototype (June 27, 2026)"],
        [
            ["1. User / Ingress", "API Gateway, UI, IAM", "Partial — CLI chat + FastAPI; UI wireframe only"],
            ["2. Agent control", "8 agents, per-agent LLM", "Complete"],
            ["3. A2A", "Message bus, remote workers", "Partial — in-process bus + file queue + journal"],
            ["4. Workflow runtime", "Steps 0–6, Clarify/Escalate", "Complete happy path; no Clarify/Escalate"],
            ["5. Tool & evidence", "MCP, vector DB, ledger", "Partial — FAISS tool + JSONL ledger"],
        ],
    )

    add_heading(doc, "6. Work Completed", 1)
    add_table(
        doc,
        ["Period / Track", "Deliverable"],
        [
            ["Phase 0", "ingest.py, FAISS, ask.py, IEEE KB, env stabilization"],
            ["Track B", "Eight agents, WorkflowEngine, typed contracts, agentic workflow"],
            ["SLM track", "Ollama integration, slm_helpers.py, parsers, heuristics, fast mode"],
            ["Track C", "Evidence Curator, Critic, evidence ledger"],
            ["Router", "Simple vs multi-hop classification"],
            ["API / A2A", "FastAPI ingress, SSE stream, MCP-style tools, A2A bus + journal"],
            ["Multi-hop hardening", "Finalize helpers, Summit/Greenfield formatters, eval tightening"],
            ["UX (local)", "Default chat mode, fast mode, consolidated docs/"],
            ["Corporate ingest", "Word, Excel, PowerPoint, OCR PDF, password PDF"],
            ["UI planning", "Professional wireframe mockup — implementation deferred"],
        ],
    )

    add_heading(doc, "7. Document Ingest Capabilities", 1)
    add_table(
        doc,
        ["Format", "Supported", "Method"],
        [
            [".pdf (text)", "Yes", "pypdf"],
            [".pdf (password)", "Yes", "Env / JSON map / CLI at ingest"],
            [".pdf (scanned)", "Yes", "OCR via Tesseract when text sparse"],
            [".docx", "Yes", "python-docx"],
            [".xlsx / .xls", "Yes", "openpyxl / xlrd"],
            [".pptx", "Yes", "python-pptx"],
            [".txt / .md", "Yes", "Direct read"],
        ],
    )
    add_para(doc, "User workflow: (1) Add files to docs/  (2) python ingest.py  (3) python ask.py")

    add_heading(doc, "8. SLM Implementation", 1)
    add_table(
        doc,
        ["Item", "Detail"],
        [
            ["Decision", "SLM-first per employer direction — not GPT/cloud default"],
            ["Model", "Meta Llama 3.2 3B via Ollama"],
            ["Verification", "LLM: ollama/llama3.2:3b @ http://localhost:11434"],
            ["Guardrails", "src/slm_helpers.py — parsers, route/plan heuristics, factual fallbacks"],
            ["Performance", "Fast mode default — fewer LLM calls, lower top-k"],
            ["Hybrid option", "Per-agent MA_RAG_<AGENT>_PROVIDER reserved for future"],
        ],
    )

    add_heading(doc, "9. Validation Results (June 27, 2026)", 1)
    add_para(doc, "pytest: 39 passed")
    add_para(doc, "eval_kb.py: 8 / 8 passed")
    add_table(
        doc,
        ["#", "Type", "Question (summary)", "Expected"],
        [
            ["1", "Single-hop", "Current completed phase", "Phase 0"],
            ["2", "Single-hop", "Volunteer researcher", "Bhargav Boyapati"],
            ["3", "Single-hop", "Stakeholder company", "Oracle"],
            ["4", "Multi-hop", "Collaborator + Oracle title", "Chandra + AI Technical Director"],
            ["5", "Single-hop", "Next phase", "Phase 1"],
            ["6", "Single-hop", "IEEE task force name", "Talent Meets AI"],
            ["7", "Single-hop", "Retrieval technology", "FAISS"],
            ["8", "Single-hop", "Research system name", "MA-RAG"],
        ],
    )
    add_heading(doc, "Recommended Health Check", 2)
    add_bullets(
        doc,
        [
            "cd ~/Projects/MA-RAG && source .venv/bin/activate",
            "python ingest.py",
            "PYTHONPATH=. pytest -q",
            "python eval_kb.py",
            'python ask.py "What is the current completed phase of the MA-RAG prototype?"',
        ],
    )

    add_heading(doc, "10. User-Facing Interfaces", 1)
    add_table(
        doc,
        ["Interface", "Status", "Access"],
        [
            ["Terminal chat", "Operational", "python ask.py"],
            ["One-shot query", "Operational", 'python ask.py "question"'],
            ["HTTP API", "Operational", "python run_api.py → port 8000"],
            ["Live agent trace", "Operational", "SSE POST /ask/stream"],
            ["Web UI", "Not built", "Wireframe prepared; deferred"],
        ],
    )

    add_heading(doc, "11. Audit and Governance", 1)
    add_table(
        doc,
        ["Artifact", "Location", "Contents"],
        [
            ["Run ID", "Printed in output", "Unique 12-char identifier"],
            ["Evidence ledger", "data/evidence_ledger/<run_id>.jsonl", "Agent steps, retrieval, conclusions"],
            ["A2A journal", "data/a2a_journal/<run_id>.jsonl", "Structured agent messages"],
        ],
    )

    add_heading(doc, "12. Known Limitations", 1)
    add_table(
        doc,
        ["Limitation", "Impact", "Planned mitigation"],
        [
            ["3B SLM quality", "Weaker on complex planning", "Heuristics; hybrid or larger SLM"],
            ["Plain-text index", "Confidential text not encrypted post-ingest", "Separate index + encryption"],
            ["No authentication", "All indexed docs queryable on machine", "API auth + RBAC"],
            ["Manual re-index", "Must re-run ingest after doc changes", "Incremental ingest"],
            ["No web UI", "Terminal/API only", "Phase 2 UI on /ask/stream"],
            ["Local uncommitted changes", "Milestone not yet in git", "Commit when approved"],
            ["Not cloud-hosted", "Single-machine prototype", "Enterprise deployment phase"],
        ],
    )

    add_heading(doc, "13. Git and Repository Status", 1)
    add_table(
        doc,
        ["Item", "Value"],
        [
            ["Branch", "feature/track-b-agentic"],
            ["Latest commit", "ea9c0e0 — API tests, A2A journal, multi-hop workflow hardening"],
            ["Local changes", "~22 modified files + new files — not yet committed"],
            ["Remote", "https://github.com/Bhargav7675/M-Rag.git"],
        ],
    )

    add_heading(doc, "14. Roadmap", 1)
    add_table(
        doc,
        ["Priority", "Item", "Timeline"],
        [
            ["P1", "Commit and stabilize current milestone", "Near term"],
            ["P2", "Web UI (chat + live agent trace sidebar)", "Next phase"],
            ["P3", "Confidential corpus segregation + redacted audit", "Security"],
            ["P4", "Broader casual-question eval suite", "Quality"],
            ["P5", "CLARIFY / ESCALATE workflow branches", "Enterprise runtime"],
            ["P6", "Production mapping (managed vector DB, queue, IAM)", "Long term"],
        ],
    )

    add_heading(doc, "15. Conclusion", 1)
    add_para(
        doc,
        "Through June 27, 2026, the MA-RAG prototype for the IEEE Talent Meets AI task force has achieved "
        "its primary research milestone: an eight-agent document-grounded Q&A pipeline with on-premises SLM "
        "(Ollama llama3.2:3b), FAISS retrieval, corporate document ingest, JSONL audit trails, API ingress "
        "with live SSE agent tracing, 39/39 automated tests, and 8/8 knowledge-base eval passing. "
        "The system is demonstration-ready on a local workstation and architecturally aligned with the "
        "enterprise five-plane multi-agent RAG reference design.",
    )

    add_heading(doc, "16. Appendix — Key Commands", 1)
    add_table(
        doc,
        ["Purpose", "Command"],
        [
            ["Index documents", "python ingest.py"],
            ["Chat", "python ask.py"],
            ["Run tests", "PYTHONPATH=. pytest -q"],
            ["KB eval", "python eval_kb.py"],
            ["Start API", "python run_api.py"],
        ],
    )

    doc.add_paragraph()
    add_para(doc, "Prepared for: Chandra Shekar Konda and IEEE Talent Meets AI stakeholders")
    add_para(doc, "Prepared by: Bhargav Boyapati, Volunteer Researcher")
    add_para(doc, "Date: June 27, 2026")

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_report()
    doc.save(str(OUTPUT))
    doc.save(str(OUTPUT_ROOT))
    print(f"Wrote {OUTPUT}")
    print(f"Wrote {OUTPUT_ROOT}")


if __name__ == "__main__":
    main()
